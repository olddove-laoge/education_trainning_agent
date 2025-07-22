import os
from docx import Document
from config import KNOWLEDGE_BASE_PATH
import re

def clean_text(text):
    """清理文本中的多余空格和换行"""
    # 合并多个空格
    text = re.sub(r'\s+', ' ', text)
    # 合并多个换行
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def read_selected_docs(selected_docs, max_length=5000):
    """读取选中的文档内容"""
    knowledge_content = ""
    
    for filename in selected_docs:
        try:
            doc_path = os.path.join(KNOWLEDGE_BASE_PATH, filename)
            if not os.path.exists(doc_path):
                continue
                
            doc = Document(doc_path)
            file_content = f"# 文档: {filename}\n\n"
            
            # 读取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    file_content += clean_text(para.text) + "\n"
            
            # 读取表格
            for table in doc.tables:
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        row_content.append(clean_text(cell.text))
                    file_content += " | ".join(row_content) + "\n"
            
            knowledge_content += file_content + "\n\n"
        except Exception as e:
            print(f"读取文档 {filename} 出错: {str(e)}")
    
    return knowledge_content[:max_length] if knowledge_content else ""